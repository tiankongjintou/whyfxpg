"""
Word report renderer: ReportRenderer port implementation using python-docx.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from whyfxpg.ports.report_renderer import ReportRenderer
from whyfxpg.services.report_model import ReportModel


class WordReportRenderer(ReportRenderer):
    """Render a ReportModel as a .docx file."""

    def render(self, report_model: ReportModel, output_path: Path) -> Path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        title = doc.add_heading("进口机电产品风险评估报告", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph(
            f"生成时间：{report_model.generated_at}"
        )
        doc.add_paragraph()

        # 一、评估概述
        doc.add_heading("一、评估概述", level=1)
        doc.add_paragraph(report_model.executive_summary)
        doc.add_paragraph(
            f"本期详细数据：共纳入 {report_model.total_events} 条风险事件，"
            "覆盖多个国家和地区。"
        )
        doc.add_paragraph("风险等级分布：")
        for level in ["S", "M", "L", "A"]:
            count = report_model.level_counts.get(level, 0)
            doc.add_paragraph(f"  {level}级风险：{count} 条", style="List Bullet")

        # 二、高风险产品清单
        doc.add_heading("二、高风险产品清单", level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        headers = ["产品名称", "品牌", "原产国", "风险等级", "综合得分", "危害类型"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for product in report_model.top_products:
            row_cells = table.add_row().cells
            row_cells[0].text = product.get("product_name", "") or "未知"
            row_cells[1].text = product.get("brand", "") or "未知"
            row_cells[2].text = product.get("country", "") or "未知"
            row_cells[3].text = product.get("latest_rs_level", "") or ""
            row_cells[4].text = str(product.get("latest_total_score", "")) or ""
            row_cells[5].text = product.get("highest_hazard_type", "") or ""

        # 三、国别风险画像
        doc.add_heading("三、国别风险画像", level=1)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Light Grid Accent 1"
        headers = ["国别", "事件总数", "S级", "M级", "L级", "A级"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for country in report_model.top_countries:
            row_cells = table.add_row().cells
            row_cells[0].text = country.get("country", "") or "未知"
            row_cells[1].text = str(country.get("event_count", 0))
            row_cells[2].text = str(country.get("s_count", 0))
            row_cells[3].text = str(country.get("m_count", 0))
            row_cells[4].text = str(country.get("l_count", 0))
            row_cells[5].text = str(country.get("a_count", 0))

        # 四、预警清单
        doc.add_heading("四、预警清单", level=1)
        for alert in report_model.pending_alerts:
            doc.add_paragraph(
                f"[{alert.get('severity', '')}] {alert.get('rule_name', '')}: "
                f"{alert.get('description', '')}",
                style="List Bullet",
            )

        # 五、方法说明
        doc.add_heading("五、评估方法说明", level=1)
        doc.add_paragraph(
            "本报告采用半定量风险矩阵法，综合考虑严重度（SS）、概率（PS）、"
            "国别修正、产品修正、历史事件密度修正和证据来源修正，"
            "最终得到风险等级S/M/L/A。"
        )

        doc.save(str(output_path))
        return output_path
